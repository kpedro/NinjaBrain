import streamlit as st
import os
from dotenv import load_dotenv
from docx import Document
import io

# Carrega variáveis do .env
load_dotenv()

# Remove BOM de variáveis de ambiente (corrige encoding UTF-8 com BOM)
for key in list(os.environ.keys()):
    if key.startswith('\ufeff'):
        new_key = key.replace('\ufeff', '')
        os.environ[new_key] = os.environ[key]
        del os.environ[key]

# 1. Configuração de Layout
st.set_page_config(
    page_title="🥷 NinjaBrain - Mentor de Vida",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 2. Configuração da API (Anthropic ou OpenAI)
def get_api_client():
    """Retorna cliente da API configurado (Anthropic ou OpenAI)"""
    # Tenta Anthropic primeiro
    anthropic_key = os.getenv("ANTHROPIC_API_KEY")
    if anthropic_key:
        try:
            import anthropic
            return anthropic.Anthropic(api_key=anthropic_key), "anthropic"
        except ImportError:
            st.warning("⚠️ Biblioteca 'anthropic' não instalada. Instale com: pip install anthropic")
    
    # Tenta OpenAI como fallback
    openai_key = os.getenv("OPENAI_API_KEY")
    if openai_key:
        try:
            import openai
            return openai.OpenAI(api_key=openai_key), "openai"
        except ImportError:
            st.warning("⚠️ Biblioteca 'openai' não instalada. Instale com: pip install openai")
    
    # Tenta Streamlit Secrets (para deploy)
    if "ANTHROPIC_API_KEY" in st.secrets:
        try:
            import anthropic
            return anthropic.Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"]), "anthropic"
        except:
            pass
    
    if "OPENAI_API_KEY" in st.secrets:
        try:
            import openai
            return openai.OpenAI(api_key=st.secrets["OPENAI_API_KEY"]), "openai"
        except:
            pass
    
    return None, None

# 3. Função para gerar resposta
def gerar_resposta(client, provider, prompt, uploaded_files=None):
    """Gera resposta usando Anthropic ou OpenAI"""
    contexto_mentor = """Você é um Mentor de Vida especializado em desenvolvimento pessoal, 
    produtividade e bem-estar. Seja empático, prático e motivador. Forneça conselhos 
    acionáveis e sempre termine com uma reflexão ou ação concreta."""
    
    mensagem_completa = f"{contexto_mentor}\n\nPergunta: {prompt}"
    
    # Processa arquivos uploadados se houver
    if uploaded_files:
        textos_arquivos = []
        for uploaded_file in uploaded_files:
            try:
                if uploaded_file.type == "text/plain":
                    texto = str(uploaded_file.read(), "utf-8")
                    textos_arquivos.append(f"\n--- Conteúdo do arquivo {uploaded_file.name} ---\n{texto}")
                elif uploaded_file.type == "text/markdown" or uploaded_file.name.endswith('.md'):
                    texto = str(uploaded_file.read(), "utf-8")
                    textos_arquivos.append(f"\n--- Conteúdo do arquivo {uploaded_file.name} ---\n{texto}")
            except Exception as e:
                st.warning(f"⚠️ Erro ao ler arquivo {uploaded_file.name}: {e}")
        
        if textos_arquivos:
            mensagem_completa += "\n\n" + "\n".join(textos_arquivos)
    
    try:
        if provider == "anthropic":
            message = client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=4096,
                messages=[{"role": "user", "content": mensagem_completa}]
            )
            return message.content[0].text
        
        elif provider == "openai":
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": mensagem_completa}],
                max_tokens=4096
            )
            return response.choices[0].message.content
        
    except Exception as e:
        raise Exception(f"Erro na API: {str(e)}")

# 4. Função para exportar em Word
def criar_docx(texto, nome_arquivo="mentoria_ninja.docx"):
    """Cria arquivo Word com o texto"""
    doc = Document()
    doc.add_heading('🥷 Mentoria NinjaBrain', 0)
    doc.add_paragraph(texto)
    
    # Salva em buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# 5. BARRA LATERAL
with st.sidebar:
    st.title("🧰 Ferramentas Ninja")
    st.success("🎯 Modo: Mentor de Vida")
    st.divider()
    
    # Upload de arquivos
    st.subheader("📥 Upload de Arquivos")
    uploaded_files = st.file_uploader(
        "Envie arquivos de texto ou Markdown",
        type=['txt', 'md'],
        accept_multiple_files=True
    )
    
    if uploaded_files:
        st.success(f"✅ {len(uploaded_files)} arquivo(s) carregado(s)")
        for file in uploaded_files:
            st.caption(f"📄 {file.name}")
    
    st.divider()
    
    # Botões de exportação (só aparecem se houver resposta)
    if "ultima_resposta" in st.session_state and st.session_state.ultima_resposta:
        st.subheader("📤 Exportar Mentoria")
        
        # Exportar TXT
        st.download_button(
            "📥 Baixar em TXT",
            st.session_state.ultima_resposta,
            file_name="mentoria_ninja.txt",
            mime="text/plain"
        )
        
        # Exportar Word
        try:
            docx_buffer = criar_docx(st.session_state.ultima_resposta)
            st.download_button(
                "📄 Salvar para Word",
                docx_buffer,
                file_name="mentoria_ninja.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Erro ao criar Word: {e}")
            st.info("Instale: pip install python-docx")

# 6. INTERFACE PRINCIPAL
st.title("🥷 NinjaBrain: Mentor de Vida")

# Verifica se há API configurada
client, provider = get_api_client()

if not client:
    st.error("""
    ❌ **Nenhuma API configurada!**
    
    Configure uma das opções no arquivo `.env`:
    - `ANTHROPIC_API_KEY=sua_chave_aqui` (recomendado)
    - `OPENAI_API_KEY=sua_chave_aqui`
    
    Ou configure nos Secrets do Streamlit Cloud.
    """)
    st.stop()

st.info(f"✅ Conectado via: **{provider.upper()}**")

# Histórico de mensagens
if "messages" not in st.session_state:
    st.session_state.messages = []

# Exibe histórico
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# Input do usuário
if prompt := st.chat_input("Digite sua pergunta ou 'oi' para começar..."):
    # Adiciona mensagem do usuário
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)
    
    # Gera resposta
    with st.chat_message("assistant"):
        with st.spinner("🧠 Mentor pensando..."):
            try:
                resposta = gerar_resposta(client, provider, prompt, uploaded_files)
                st.markdown(resposta)
                
                # Salva resposta para exportação
                st.session_state.ultima_resposta = resposta
                
                # Adiciona ao histórico
                st.session_state.messages.append({"role": "assistant", "content": resposta})
                
            except Exception as e:
                erro_msg = f"❌ Erro: {str(e)}"
                st.error(erro_msg)
                st.session_state.messages.append({"role": "assistant", "content": erro_msg})
